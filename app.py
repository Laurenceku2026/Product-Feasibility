import streamlit as st
import openai
import json
import os
import re
import secrets
import string
import time
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.header import Header
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime, timedelta

# ================== 页面配置 ==================
st.set_page_config(
    page_title="Product Feasibility Analysis System",
    page_icon="📊",
    layout="wide"
)

# ================== 管理员凭证 ==================
ADMIN_USERNAME = "Laurence_ku"
ADMIN_PASSWORD = "Ku_product$2026"

# ================== 从 secrets 读取永久 API 配置 ==================
try:
    PERSISTENT_API_KEY = st.secrets["AI_API_KEY"]
except:
    PERSISTENT_API_KEY = ""
try:
    PERSISTENT_BASE_URL = st.secrets["AI_BASE_URL"]
except:
    PERSISTENT_BASE_URL = "https://api.deepseek.com"
try:
    PERSISTENT_MODEL_NAME = st.secrets["AI_MODEL_NAME"]
except:
    PERSISTENT_MODEL_NAME = "deepseek-chat"

# ================== 从 secrets 读取 SMTP 邮件配置 ==================
try:
    SMTP_SERVER = st.secrets["SMTP_SERVER"]
    SMTP_PORT = st.secrets["SMTP_PORT"]
    SMTP_USER = st.secrets["SMTP_USER"]
    SMTP_PASSWORD = st.secrets["SMTP_PASSWORD"]
except:
    SMTP_SERVER = ""
    SMTP_PORT = 587
    SMTP_USER = ""
    SMTP_PASSWORD = ""

def send_license_email(to_email, license_key, plan_name, uses, expiry, lang="zh"):
    """发送授权码邮件，支持中英文"""
    if not SMTP_USER or not SMTP_PASSWORD:
        return False, "邮件服务未配置"
    
    if lang == "zh":
        subject = f"您的产品可行性分析系统授权码 - {plan_name}"
        body = f"""
亲爱的用户：

感谢您对 Techlife 产品的信任！

您的产品分析报告通行证已生成，详情如下：

- 授权码：{license_key}
- 套餐：{plan_name}
- 可用次数：{uses}
- 有效期至：{expiry}

请在系统左侧边栏的“授权码 (Report Key)”输入框中输入此授权码，即可解锁高级功能（无水印、可下载 Word 报告）。

请妥善保管此授权码，如有疑问请联系：nc.ku@hotmail.com

祝您使用愉快！

Techlife 产品可行性分析系统
"""
    else:
        subject = f"Your License Key for Product Feasibility Analysis System - {plan_name}"
        body = f"""
Dear Customer,

Thank you for trusting Techlife products!

Your product analysis report pass has been generated. Details are as follows:

- License Key: {license_key}
- Plan: {plan_name}
- Available uses: {uses}
- Valid until: {expiry}

Please enter this license key in the "Report Key" input box on the left sidebar to unlock advanced features (no watermark, Word report download available).

Please keep this license key safe. If you have any questions, please contact: nc.ku@hotmail.com

Best regards,

Techlife Product Feasibility Analysis System
"""
    msg = MIMEMultipart()
    msg['From'] = SMTP_USER
    msg['To'] = to_email
    msg['Subject'] = Header(subject, 'utf-8')
    msg.attach(MIMEText(body, 'plain', 'utf-8'))
    try:
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
        server.starttls()
        server.login(SMTP_USER, SMTP_PASSWORD)
        server.sendmail(SMTP_USER, [to_email], msg.as_string())
        server.quit()
        return True, ""
    except Exception as e:
        return False, str(e)

# ================== 授权类型定义 ==================
LICENSE_TYPES = {
    "trial": {"name": "试用版", "max_uses": 60, "max_months": 3, "en_name": "Trial"},
    "level1": {"name": "一级用户", "max_uses": 100, "max_months": 12, "en_name": "Level 1"},
    "level2": {"name": "二级用户", "max_uses": 300, "max_months": 24, "en_name": "Level 2"},
    "level3": {"name": "三级用户", "max_uses": 500, "max_months": 36, "en_name": "Level 3"},
    "level4": {"name": "四级用户", "max_uses": 1000, "max_months": 60, "en_name": "Level 4"},
}

USAGE_FILE = "usage_data.json"

def load_usage_data():
    if os.path.exists(USAGE_FILE):
        try:
            with open(USAGE_FILE, "r") as f:
                return json.load(f)
        except:
            return {}
    return {}

def save_usage_data(data):
    with open(USAGE_FILE, "w") as f:
        json.dump(data, f, indent=2)

# ================== 初始化 session state (新增浏览器语言检测) ==================
# 检测浏览器语言并初始化默认语言
if "lang" not in st.session_state:
    # 注入JS获取浏览器语言，并通过query params传递
    st.markdown("""
    <script>
    // 获取浏览器语言
    const browserLang = navigator.language || navigator.userLanguage;
    // 只保留前两位（如zh-CN -> zh, en-US -> en）
    const shortLang = browserLang.substring(0, 2);
    // 如果当前URL没有lang参数，自动添加并刷新
    const urlParams = new URLSearchParams(window.location.search);
    if (!urlParams.has('lang')) {
        urlParams.set('lang', shortLang);
        window.location.search = urlParams;
    }
    </script>
    """, unsafe_allow_html=True)
    
    # 从query params获取语言，默认英文
    query_params = st.query_params
    browser_lang = query_params.get("lang", "en")
    # 映射语言：中文相关（zh）→ zh，其他→ en
    st.session_state.lang = "zh" if browser_lang.lower() == "zh" else "en"

if "pulse_active" not in st.session_state:
    st.session_state.pulse_active = False
if "report_content_zh" not in st.session_state:
    st.session_state.report_content_zh = None
if "report_content_en" not in st.session_state:
    st.session_state.report_content_en = None
if "admin_logged_in" not in st.session_state:
    st.session_state.admin_logged_in = False
if "ai_api_key" not in st.session_state:
    st.session_state.ai_api_key = PERSISTENT_API_KEY
if "ai_base_url" not in st.session_state:
    st.session_state.ai_base_url = PERSISTENT_BASE_URL
if "ai_model_name" not in st.session_state:
    st.session_state.ai_model_name = PERSISTENT_MODEL_NAME
if "usage_db" not in st.session_state:
    st.session_state.usage_db = load_usage_data()
if "current_report_key" not in st.session_state:
    st.session_state.current_report_key = ""
if "current_license_type" not in st.session_state:
    st.session_state.current_license_type = None

def activate_license(report_key):
    if report_key in st.session_state.usage_db:
        record = st.session_state.usage_db[report_key]
        remaining = record["remaining"]
        expiry_str = record["expiry"]
        expiry = datetime.fromisoformat(expiry_str)
        if remaining > 0 and datetime.now() <= expiry:
            lic_type = record.get("type", "unknown")
            st.session_state.current_license_type = lic_type
            return True, remaining, expiry_str, lic_type
        else:
            st.session_state.current_license_type = None
            return False, 0, None, None
    else:
        st.session_state.current_license_type = None
        return False, 0, None, None

def consume_usage(report_key):
    if st.session_state.admin_logged_in:
        return True
    if not report_key:
        return False
    valid, remaining, expiry_str, _ = activate_license(report_key)
    if not valid:
        return False
    record = st.session_state.usage_db[report_key]
    record["remaining"] -= 1
    record["total_uses"] = record.get("total_uses", 0) + 1
    save_usage_data(st.session_state.usage_db)
    return True

def get_remaining_info(report_key):
    if st.session_state.admin_logged_in:
        return "无限", "永久"
    valid, remaining, expiry_str, _ = activate_license(report_key)
    if not valid:
        return "未授权", "无"
    expiry = datetime.fromisoformat(expiry_str)
    return str(remaining), expiry.strftime("%Y-%m-%d")

def is_premium_user(report_key):
    if st.session_state.admin_logged_in:
        return True
    if report_key:
        valid, _, _, _ = activate_license(report_key)
        return valid
    return False

def generate_report_key(license_type, custom_uses=None, custom_months=None):
    random_str = ''.join(secrets.choice(string.ascii_uppercase + string.digits) for _ in range(8))
    new_key = f"{license_type.upper()}_{random_str}"
    if license_type == "custom":
        max_uses = custom_uses
        max_months = custom_months
        type_name = "自定义"
    else:
        lic_info = LICENSE_TYPES[license_type]
        max_uses = lic_info["max_uses"]
        max_months = lic_info["max_months"]
        type_name = lic_info["name"]
    expiry = datetime.now() + timedelta(days=max_months*30)
    expiry_str = expiry.isoformat()
    st.session_state.usage_db[new_key] = {
        "type": license_type,
        "remaining": max_uses,
        "expiry": expiry_str,
        "total_uses": 0,
        "generated_at": datetime.now().isoformat()
    }
    save_usage_data(st.session_state.usage_db)
    return new_key, max_uses, expiry_str, type_name

# ================== 防复制/截屏 CSS + 动态水印 ==================
def add_security_css(disable=False):
    if disable:
        return
    st.markdown("""
    <style>
        body, .stApp, .markdown-text-container, .report-container {
            user-select: none;
            -webkit-user-select: none;
            -moz-user-select: none;
            -ms-user-select: none;
        }
        body {
            -webkit-touch-callout: none;
            pointer-events: auto;
        }
        .bg-watermark {
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
            height: 100%;
            opacity: 0.05;
            pointer-events: none;
            z-index: 999;
            background-image: repeating-linear-gradient(45deg, 
                #000 0px, #000 2px, 
                transparent 2px, transparent 40px,
                #000 40px, #000 42px,
                transparent 42px, transparent 80px);
            background-size: 80px 80px;
        }
    </style>
    <div class="bg-watermark"></div>
    <script>
        document.addEventListener('contextmenu', function(e) {
            e.preventDefault();
            return false;
        });
        document.addEventListener('keydown', function(e) {
            if (e.ctrlKey && (e.key === 'c' || e.key === 'C' || e.key === 'v' || e.key === 'V' || e.key === 'x' || e.key === 'X' || e.key === 's' || e.key === 'S')) {
                e.preventDefault();
                return false;
            }
            if (e.key === 'F12' || e.key === 'PrintScreen') {
                e.preventDefault();
                alert('截图功能已被禁用，请遵守保密协议。');
                return false;
            }
        });
        document.addEventListener('keyup', function(e) {
            if (e.key === 'PrintScreen') {
                alert('截图行为已被记录，请勿传播保密内容。');
            }
        });
    </script>
    """, unsafe_allow_html=True)

def add_dynamic_watermark(lang, hide):
    if hide:
        return
    if lang == "zh":
        watermark_text = "机密，样板报告，请联系 nc.ku@hotmail.com"
    else:
        watermark_text = "Confidential, Sample Report, Pls contact nc.ku@hotmail.com"
    st.markdown(f"""
    <div style="position: fixed; bottom: 20px; right: 20px; opacity: 0.4; font-size: 14px; color: #666; pointer-events: none; z-index: 1000; font-family: sans-serif; background: rgba(255,255,255,0.5); padding: 4px 8px; border-radius: 4px;">
        {watermark_text}
    </div>
    """, unsafe_allow_html=True)

# ================== Word 表格生成（自动列宽，浅灰边框） ==================
def set_cell_border(cell, border_color=RGBColor(0xCC, 0xCC, 0xCC)):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ['top', 'left', 'bottom', 'right']:
        tag = f'w:{edge}'
        border = OxmlElement(tag)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), f'{border_color}')
        tcPr.append(border)

def markdown_to_docx(md_text, doc, lang):
    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
            i += 1
            continue
        if line.startswith('## '):
            doc.add_heading(line[3:], level=2)
            i += 1
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
            i += 1
            continue
        if line.startswith('|') and i+1 < len(lines):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            if len(table_lines) >= 2:
                def parse_row(row):
                    cells = [cell.strip() for cell in row.split('|')]
                    if cells and cells[0] == '':
                        cells = cells[1:]
                    if cells and cells[-1] == '':
                        cells = cells[:-1]
                    return cells
                headers = parse_row(table_lines[0])
                if len(table_lines) > 1 and '---' in table_lines[1]:
                    data_lines = table_lines[2:] if len(table_lines) > 2 else []
                else:
                    data_lines = table_lines[1:]
                num_cols = len(headers)
                if num_cols > 0:
                    table = doc.add_table(rows=1+len(data_lines), cols=num_cols)
                    table.style = 'Table Grid'
                    table.autofit = True
                    table.width = Inches(6.5)
                    for row in table.rows:
                        for cell in row.cells:
                            set_cell_border(cell, RGBColor(0xCC, 0xCC, 0xCC))
                    for col_idx, cell_text in enumerate(headers):
                        cell = table.cell(0, col_idx)
                        cell.text = cell_text
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.name = 'Arial' if lang == 'en' else '宋体'
                    for row_idx, data_line in enumerate(data_lines):
                        cells = parse_row(data_line)
                        for col_idx, cell_text in enumerate(cells):
                            if col_idx < num_cols:
                                cell = table.cell(row_idx+1, col_idx)
                                cell.text = cell_text
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.name = 'Arial' if lang == 'en' else '宋体'
                    doc.add_paragraph()
            continue
        if line.strip():
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.name = 'Arial' if lang == 'en' else '宋体'
        else:
            doc.add_paragraph()
        i += 1

# ================== 管理员对话框 ==================
@st.dialog("管理员登录")
def admin_login_dialog():
    username = st.text_input("用户名")
    password = st.text_input("密码", type="password")
    if st.button("登录"):
        if username == ADMIN_USERNAME and password == ADMIN_PASSWORD:
            st.session_state.admin_logged_in = True
            st.success("登录成功！")
            st.rerun()
        else:
            st.error("用户名或密码错误")

@st.dialog("管理员设置")
def admin_settings_dialog():
    st.subheader("AI API 配置（临时覆盖）")
    new_key = st.text_input("API Key", value=st.session_state.ai_api_key, type="password")
    new_url = st.text_input("Base URL", value=st.session_state.ai_base_url)
    new_model = st.text_input("模型名称", value=st.session_state.ai_model_name)
    if st.button("应用临时配置"):
        st.session_state.ai_api_key = new_key
        st.session_state.ai_base_url = new_url
        st.session_state.ai_model_name = new_model
        st.success("当前会话已使用新配置（刷新页面后恢复为永久配置）")
        st.rerun()
    st.markdown("---")
    
    st.subheader("Report Key 生成器")
    key_type = st.selectbox("选择授权类型", ["试用版", "一级用户", "二级用户", "三级用户", "四级用户", "自定义"])
    custom_uses = None
    custom_months = None
    if key_type == "自定义":
        col_c1, col_c2 = st.columns(2)
        with col_c1:
            custom_uses = st.number_input("使用次数", min_value=1, step=1, value=100)
        with col_c2:
            custom_months = st.number_input("有效期（月）", min_value=1, step=1, value=12)
    if st.button("生成 Report Key"):
        if key_type == "试用版":
            license_type = "trial"
        elif key_type == "一级用户":
            license_type = "level1"
        elif key_type == "二级用户":
            license_type = "level2"
        elif key_type == "三级用户":
            license_type = "level3"
        elif key_type == "四级用户":
            license_type = "level4"
        else:
            license_type = "custom"
        new_key, max_uses, expiry_str, type_name = generate_report_key(license_type, custom_uses, custom_months)
        st.success(f"已生成 {type_name} Report Key：")
        st.code(new_key, language="text")
        st.write(f"可使用次数：{max_uses} 次，有效期至：{expiry_str[:10]}")
    
    st.markdown("---")
    st.subheader("生成付费套餐授权码")

    col_price1, col_price2, col_price3 = st.columns(3)
    with col_price1:
        st.markdown("**单次通行**")
        st.markdown("18元 / 3美元")
        st.markdown("1次 · 无有效期")
        if st.button("生成单次通行码"):
            new_key, max_uses, expiry_str, _ = generate_report_key("custom", custom_uses=1, custom_months=9999)
            st.success(f"单次通行授权码：")
            st.code(new_key, language="text")
            st.write(f"次数：{max_uses}，有效期：无限制（至 {expiry_str[:10]}）")
    
    with col_price2:
        st.markdown("**100次套餐**")
        st.markdown("180元 / 30美元")
        st.markdown("100次 · 1个月")
        if st.button("生成100次套餐码"):
            new_key, max_uses, expiry_str, _ = generate_report_key("custom", custom_uses=100, custom_months=1)
            st.success(f"100次套餐授权码：")
            st.code(new_key, language="text")
            st.write(f"次数：{max_uses}，有效期：1个月（至 {expiry_str[:10]}）")
    
    with col_price3:
        st.markdown("**1200次套餐**")
        st.markdown("1200元 / 200美元")
        st.markdown("1200次 · 12个月")
        if st.button("生成1200次套餐码"):
            new_key, max_uses, expiry_str, _ = generate_report_key("custom", custom_uses=1200, custom_months=12)
            st.success(f"1200次套餐授权码：")
            st.code(new_key, language="text")
            st.write(f"次数：{max_uses}，有效期：12个月（至 {expiry_str[:10]}）")
    
    st.markdown("---")
    st.subheader("已生成的所有 Report Key")
    for key, data in st.session_state.usage_db.items():
        expiry = datetime.fromisoformat(data["expiry"]).strftime("%Y-%m-%d")
        st.write(f"- {key}: {data['remaining']} 次剩余, 有效期至 {expiry}")
    
    st.markdown("---")
    st.subheader("永久修改 API Key")
    st.markdown("请前往 [Streamlit Cloud Secrets](https://share.streamlit.io/) 修改 `AI_API_KEY`、`AI_BASE_URL` 和 `AI_MODEL_NAME`，然后重启应用。")

# ================== 右上角按钮 ==================
col1, col2, col3, col4 = st.columns([8, 1, 1, 1])
with col2:
    if st.button("中文", key="zh_btn"):
        st.session_state.lang = "zh"
        st.rerun()
with col3:
    if st.button("English", key="en_btn"):
        st.session_state.lang = "en"
        st.rerun()
with col4:
    if st.button("⚙️", key="settings_btn"):
        if st.session_state.admin_logged_in:
            admin_settings_dialog()
        else:
            admin_login_dialog()

# ================== 语言文本（包含完整 report_prompt） ==================
TEXTS = {
    "zh": {
        "title": "📊 产品可行性 - AI分析系统",
        "sidebar_title": "关于分析系统",
        "sidebar_basis": "本系统基于：",
        "basis_items": ["25+年研发管理经验", "AI大模型数据分析", "行业数据库与竞品追踪", "DFSS/六西格玛方法论"],
        "analyst_name_label": "分析人姓名",
        "analyst_name_ph": "请输入您的姓名或分析师姓名",
        "analyst_title_label": "分析人头衔（可选）",
        "analyst_title_ph": "例如：研发总监、技术顾问",
        "api_status": "AI API 状态",
        "api_configured": "✅ 已配置",
        "api_not_configured": "❌ 未配置，请联系管理员",
        "report_key_label": "授权码 (Report Key)",
        "report_key_help": "输入授权码可获得完整权限",
        "license_info": "授权信息",
        "remaining_label": "剩余次数",
        "expiry_label": "有效期至",
        "contact_info": "📞 **联系人：古生**  \n✉️ 电邮: nc.ku@hotmail.com  \n📱 电话/微信: +86-13823760640",
        "purchase_title": "💰 购买+解锁",
        "input_title": "📝 产品信息输入",
        "basic_info": "基本信息",
        "product_name": "产品名称",
        "product_name_ph": "例如：宠物智能饮水机",
        "product_desc": "产品简要描述",
        "product_desc_ph": "例如：一款支持APP控制、可记录饮水量的智能宠物饮水机",
        "target_markets": "目标市场",
        "market_options": ["中国大陆", "美国", "欧洲", "东南亚", "日本", "其他"],
        "target_users": "目标用户群体",
        "target_users_ph": "例如：25-40岁城市中产、养猫人群",
        "market_channel": "市场与渠道",
        "channel_status": "现有渠道情况",
        "channel_options": ["有成熟渠道", "有部分渠道", "渠道较弱", "无渠道/从零开始"],
        "channel_detail": "渠道详情",
        "channel_detail_ph": "例如：天猫旗舰店、京东自营、部分线下宠物店",
        "brand_status": "品牌认知度",
        "brand_options": ["高（知名品牌）", "中（行业内有认知）", "低（需要建立品牌）"],
        "tech_capability": "技术能力",
        "tech_experience": "相关技术经验",
        "tech_options": ["智能硬件/物联网", "APP开发", "机械结构设计", "光学设计", "电子电路", "供应链管理", "海外认证（UL/CE/FCC）", "DFSS/六西格玛"],
        "dev_stage": "产品开发阶段",
        "stage_options": ["概念/想法", "调研中", "已立项", "开发中", "已有样机"],
        "business_goals": "商业目标",
        "estimated_budget": "预估研发预算",
        "budget_options": ["50万以下", "50-100万", "100-200万", "200-500万", "500万以上"],
        "sales_target": "首年销售目标",
        "sales_target_ph": "例如：1000万人民币 / 200万美元",
        "other_info": "其他信息",
        "other_ph": "任何你认为重要的信息，如：已有技术储备、合作伙伴、特殊要求等",
        "submit_btn": "🚀 开始分析",
        "product_name_missing": "请填写产品名称",
        "api_key_missing": "AI API Key 未配置，请联系管理员",
        "generating": "报告生成中，请稍候...",
        "error_prefix": "报告生成失败：",
        "report_title": "📄 生成的可行性分析报告",
        "download_section": "📥 下载报告",
        "download_btn": "下载 Word 文档",
        "download_unlock_btn": "📥 下载报告+解锁",
        "key_error": "授权码无效或已过期",
        "back_btn": "← 返回重新填写",
        "footer": "© 2026 Laurence Ku | AI产品可行性分析系统 | 基于25年研发管理经验",
        "trial_ended": "试用已结束，请联系 nc.ku@hotmail.com",
        "no_license": "未输入授权码，当前为试用模式（有水印、不可复制、不可下载）",
        "report_prompt": """
你是一位资深产品分析师和研发顾问，拥有25年消费电子及智能硬件行业经验。请根据以下产品信息，生成一份专业的《产品可行性分析报告》。

报告必须严格按照以下Markdown结构输出，内容要具体、有洞察，数据基于行业常识合理推断。重要要求：
1. 表格必须使用标准Markdown表格语法，即使用竖线分隔单元格，第二行为分隔行（例如 |---|---|）。
2. 禁止在表格内外使用任何加粗标记（如 ** 或 *），也不要使用斜体。所有文本保持纯文本格式。
3. 禁止在表格单元格内使用换行符或复杂格式。

# 《产品可行性分析报告》
## {product_name}

## 报告基本信息

| 项目 | 内容 |
|------|------|
| 产品名称 | {product_name} |
| 产品描述 | {product_description} |
| 目标市场 | {target_markets} |
| 目标用户 | {target_users} |
| 报告日期 | {{CURRENT_DATE}} |
| 分析人 | {{ANALYST_INFO}} |

---

## 第一部分：市场需求分析

### 1.1 市场规模与趋势

（请根据目标市场分别列出主要市场的规模、增长率、驱动因素和瓶颈，用表格形式）

### 1.2 用户画像

（用表格描述核心用户特征）

### 1.3 用户痛点分析

（列出3-5个核心痛点，用表格说明提及频率和描述）

### 1.4 关键功能需求排序

（用表格列出功能、重要性评分和说明）

---

## 第二部分：竞品分析

### 2.1 主要竞争对手

（根据产品品类，列出至少3个主要竞品，用表格说明品牌、产品、优势、劣势、定价区间）

### 2.2 竞品功能对比

（选择5-6个关键功能进行对比，用表格展示）

### 2.3 市场空白点分析

（列出至少3个市场空白机会）

---

## 第三部分：渠道适配性分析

### 3.1 目标市场渠道结构

（用表格描述主要渠道类型、占比、特点、适合度）

### 3.2 客户现有渠道现状

（基于用户输入：渠道情况={channel_status}，渠道详情={channel_detail}，品牌认知度={brand_status}，进行分析）

### 3.3 渠道策略建议

（按年份给出渠道拓展建议，用表格）

---

## 第四部分：技术可行性评估

### 4.1 关键技术要求

（用表格列出关键技术项、要求、客户现有能力、风险评估）

### 4.2 开发周期估算

（用表格列出阶段、时间、关键任务）

### 4.3 关键风险点

（用表格列出风险、可能性、影响、应对措施）

---

## 第五部分：销售预测

### 5.1 预测模型假设

（列出定价、目标市场、份额等假设）

### 5.2 销售额预测

（3年预测，用表格）

### 5.3 投资回报估算

（用表格列出研发投入、市场推广、首批生产成本、总启动资金、毛利率、盈亏平衡点）

---

## 第六部分：结论与建议

### 6.1 综合评估

（用表格打分：市场吸引力、技术可行性、渠道匹配度、竞争格局、投资回报，各1-10分，并说明）

### 6.2 差异化定位建议

（给出2-3个定位选项，用表格分析优势和风险）

### 6.3 最终建议

（给出综合评分和建议的下一步行动，5点以内）

---

请直接输出报告内容，不要添加额外解释。对于用户未提供的信息，基于行业标准进行合理推断，并注明“基于行业分析”。
"""
    },
    "en": {
        "title": "📊 Product Feasibility - AI Analysis System",
        "sidebar_title": "About the System",
        "sidebar_basis": "This system is based on:",
        "basis_items": ["25+ years R&D management", "AI big data analysis", "Industry database & competitor tracking", "DFSS/Six Sigma methodology"],
        "analyst_name_label": "Analyst Name",
        "analyst_name_ph": "Enter your name or analyst name",
        "analyst_title_label": "Analyst Title (Optional)",
        "analyst_title_ph": "e.g., R&D Director, Technical Consultant",
        "api_status": "AI API Status",
        "api_configured": "✅ Configured",
        "api_not_configured": "❌ Not configured, contact admin",
        "report_key_label": "Report Key",
        "report_key_help": "Enter the license key to get full access",
        "license_info": "License Info",
        "remaining_label": "Remaining uses",
        "expiry_label": "Valid until",
        "contact_info": "📞 **Contact: Laurence Ku**  \n✉️ Email: nc.ku@hotmail.com  \n📱 Phone/Wechat: +86-13823760640",
        "purchase_title": "💰 Purchase + Unlock",
        "input_title": "📝 Product Information Input",
        "basic_info": "Basic Information",
        "product_name": "Product Name",
        "product_name_ph": "e.g., Smart Pet Water Fountain",
        "product_desc": "Brief Description",
        "product_desc_ph": "e.g., A smart pet fountain with APP control and water intake logging",
        "target_markets": "Target Markets",
        "market_options": ["Mainland China", "United States", "Europe", "Southeast Asia", "Japan", "Others"],
        "target_users": "Target User Group",
        "target_users_ph": "e.g., Urban middle-class cat owners aged 25-40",
        "market_channel": "Market & Channel",
        "channel_status": "Current Channel Status",
        "channel_options": ["Mature channels", "Partial channels", "Weak channels", "No channels / start from scratch"],
        "channel_detail": "Channel Details",
        "channel_detail_ph": "e.g., Tmall flagship store, JD self-operated, some offline pet stores",
        "brand_status": "Brand Awareness",
        "brand_options": ["High (well-known)", "Medium (recognized in industry)", "Low (need to build brand)"],
        "tech_capability": "Technical Capability",
        "tech_experience": "Relevant Tech Experience",
        "tech_options": ["Smart Hardware/IoT", "App Development", "Mechanical Design", "Optical Design", "Electronic Circuits", "Supply Chain Management", "Overseas Certification (UL/CE/FCC)", "DFSS/Six Sigma"],
        "dev_stage": "Development Stage",
        "stage_options": ["Idea/Concept", "Researching", "Project approved", "Developing", "Prototype ready"],
        "business_goals": "Business Goals",
        "estimated_budget": "Estimated R&D Budget",
        "budget_options": ["Under 500k", "500k-1M", "1M-2M", "2M-5M", "Above 5M"],
        "sales_target": "First Year Sales Target",
        "sales_target_ph": "e.g., 10M RMB / 2M USD",
        "other_info": "Other Information",
        "other_ph": "Any important info, e.g., existing tech stack, partners, special requirements",
        "submit_btn": "🚀 Start Analysis",
        "product_name_missing": "Please enter the product name",
        "api_key_missing": "AI API Key not configured, contact admin",
        "generating": "Generating report, please wait...",
        "error_prefix": "Report generation failed: ",
        "report_title": "📄 Generated Feasibility Analysis Report",
        "download_section": "📥 Download Report",
        "download_btn": "Download Word Document",
        "download_unlock_btn": "📥 Download Report + Unlock",
        "key_error": "Invalid or expired Report Key",
        "back_btn": "← Back to re-enter",
        "footer": "© 2026 Laurence Ku | AI Product Feasibility System | Based on 25+ years R&D experience",
        "trial_ended": "Trial finished, please contact nc.ku@hotmail.com",
        "no_license": "No Report Key entered. Trial mode (watermark, no copy, no download).",
        "report_prompt": """
You are a senior product analyst and R&D consultant with 25 years of experience in consumer electronics and smart hardware. Based on the following product information, generate a professional "Product Feasibility Analysis Report".

The report must strictly follow the Markdown structure below. The content should be specific, insightful, and based on industry common sense. Important requirements:
1. Tables must use standard Markdown table syntax (e.g., | Header | Header |, |---|---|).
2. Do not use any bold or italic markers (like ** or *) inside or outside tables. Keep all text plain.
3. Do not use line breaks or complex formatting inside table cells.

# Product Feasibility Analysis Report
## {product_name}

## Report Basic Information

| Item | Content |
|------|---------|
| Product Name | {product_name} |
| Product Description | {product_description} |
| Target Markets | {target_markets} |
| Target Users | {target_users} |
| Report Date | {{CURRENT_DATE}} |
| Analyst | {{ANALYST_INFO}} |

---

## Part 1: Market Demand Analysis

### 1.1 Market Size & Trends

(For each target market, list market size, growth rate, key drivers and barriers in a table)

### 1.2 User Persona

(Describe core user characteristics in a table)

### 1.3 User Pain Points

(List 3-5 core pain points in a table with frequency and description)

### 1.4 Key Feature Priority

(List features, importance score, and explanation in a table)

---

## Part 2: Competitive Analysis

### 2.1 Main Competitors

(List at least 3 main competitors with brand, product, strengths, weaknesses, price range in a table)

### 2.2 Feature Comparison

(Compare 5-6 key features in a table)

### 2.3 Market Gap Summary

(List at least 3 market gap opportunities)

---

## Part 3: Channel Suitability Analysis

### 3.1 Target Market Channel Structure

(Describe main channel types, share, characteristics, suitability in a table)

### 3.2 Client's Current Channel Status

(Based on user input: channel status={channel_status}, channel details={channel_detail}, brand awareness={brand_status})

### 3.3 Channel Strategy Recommendations

(Provide channel expansion recommendations by year in a table)

---

## Part 4: Technical Feasibility Assessment

### 4.1 Key Technical Requirements

(List technology, requirement, client capability, risk level in a table)

### 4.2 Development Timeline Estimate

(List phase, duration, key tasks in a table)

### 4.3 Key Risk Points

(List risk, probability, impact, mitigation in a table)

---

## Part 5: Sales Forecast

### 5.1 Forecast Assumptions

(List pricing, target market, share assumptions)

### 5.2 Sales Forecast

(3-year forecast in a table)

### 5.3 ROI Estimate

(List R&D investment, marketing, first production cost, total capital, gross margin, breakeven point in a table)

---

## Part 6: Conclusion & Recommendations

### 6.1 Comprehensive Evaluation

(Score each dimension: Market Attractiveness, Technical Feasibility, Channel Fit, Competitive Landscape, ROI Potential out of 10, with explanation in a table)

### 6.2 Differentiation Positioning Recommendations

(Provide 2-3 positioning options with advantages and risks in a table)

### 6.3 Final Recommendation

(Provide overall score and 5 specific next steps)

---

Output the report directly without additional explanation. For information not provided by the user, make reasonable inferences based on industry standards and note "based on industry analysis".
"""
    }
}

# ================== 获取当前语言 ==================
lang = st.session_state.lang
t = TEXTS[lang]

st.title(t["title"])

# ================== 支付回调处理 ==================
params = st.query_params
if "order_success" in params and "plan" in params:
    plan = params["plan"]
    customer_email = params.get("email", None)
    current_lang = st.session_state.lang
    
    if current_lang == "zh":
        if plan == "single":
            uses = 1
            months = 9999
            plan_name = "单次通行"
        elif plan == "100":
            uses = 100
            months = 1
            plan_name = "100次套餐"
        elif plan == "1200":
            uses = 1200
            months = 12
            plan_name = "1200次套餐"
        else:
            uses = 0
            months = 0
            plan_name = "未知"
    else:
        if plan == "single":
            uses = 1
            months = 9999
            plan_name = "Single Pass"
        elif plan == "100":
            uses = 100
            months = 1
            plan_name = "100 Credits"
        elif plan == "1200":
            uses = 1200
            months = 12
            plan_name = "1200 Credits"
        else:
            uses = 0
            months = 0
            plan_name = "Unknown"
    
    if uses > 0:
        # 生成授权码
        new_key, max_uses, expiry_str, _ = generate_report_key("custom", custom_uses=uses, custom_months=months)
        expiry_date = datetime.fromisoformat(expiry_str).strftime("%Y-%m-%d") if months != 9999 else "永久"
        
        # 发送邮件
        if customer_email:
            success, error = send_license_email(customer_email, new_key, plan_name, uses, expiry_date, current_lang)
            if success:
                st.success(f"{t['purchase_title']} - {t['api_configured'] if current_lang == 'zh' else 'Success'}!")
                st.write(f"{t['report_key_label']}: {new_key}")
                st.write(f"{t['remaining_label']}: {uses}")
                st.write(f"{t['expiry_label']}: {expiry_date}")
                st.write(f"{t['contact_info']}")
            else:
                st.error(f"{t['error_prefix']} {error}")
        else:
            st.success(f"{t['purchase_title']} - {t['api_configured'] if current_lang == 'zh' else 'Success'}!")
            st.write(f"{t['report_key_label']}: {new_key}")
            st.write(f"{t['remaining_label']}: {uses}")
            st.write(f"{t['expiry_label']}: {expiry_date}")

# ================== 侧边栏配置 ==================
with st.sidebar:
    st.header(t["sidebar_title"])
    st.write(t["sidebar_basis"])
    for item in t["basis_items"]:
        st.write(f"• {item}")
    
    st.markdown("---")
    
    # API状态显示
    st.subheader(t["api_status"])
    if st.session_state.ai_api_key:
        st.write(t["api_configured"])
    else:
        st.write(t["api_not_configured"])
    
    st.markdown("---")
    
    # 授权码输入
    st.subheader(t["report_key_label"])
    report_key = st.text_input("", placeholder=t["report_key_help"], key="report_key_input")
    if report_key:
        st.session_state.current_report_key = report_key
        valid, remaining, expiry_str, lic_type = activate_license(report_key)
        if valid:
            st.success(t["api_configured"])
            st.write(f"{t['remaining_label']}: {remaining}")
            st.write(f"{t['expiry_label']}: {expiry_str[:10]}")
        else:
            st.error(t["key_error"])
    else:
        st.info(t["no_license"])
    
    st.markdown("---")
    
    # 联系人信息
    st.markdown(t["contact_info"])

# ================== 主界面输入表单 ==================
st.subheader(t["input_title"])

with st.form(key="product_form"):
    col1, col2 = st.columns(2)
    with col1:
        analyst_name = st.text_input(t["analyst_name_label"], placeholder=t["analyst_name_ph"])
        analyst_title = st.text_input(t["analyst_title_label"], placeholder=t["analyst_title_ph"])
        product_name = st.text_input(t["product_name"], placeholder=t["product_name_ph"], required=True)
        product_desc = st.text_area(t["product_desc"], placeholder=t["product_desc_ph"], height=100)
        target_markets = st.multiselect(t["target_markets"], t["market_options"])
        target_users = st.text_input(t["target_users"], placeholder=t["target_users_ph"])
    
    with col2:
        channel_status = st.selectbox(t["channel_status"], t["channel_options"])
        channel_detail = st.text_input(t["channel_detail"], placeholder=t["channel_detail_ph"])
        brand_status = st.selectbox(t["brand_status"], t["brand_options"])
        tech_experience = st.multiselect(t["tech_experience"], t["tech_options"])
        dev_stage = st.selectbox(t["dev_stage"], t["stage_options"])
        estimated_budget = st.selectbox(t["estimated_budget"], t["budget_options"])
        sales_target = st.text_input(t["sales_target"], placeholder=t["sales_target_ph"])
        other_info = st.text_area(t["other_info"], placeholder=t["other_ph"], height=100)
    
    submit_btn = st.form_submit_button(t["submit_btn"])

# ================== 报告生成逻辑 ==================
if submit_btn:
    if not product_name:
        st.error(t["product_name_missing"])
    elif not st.session_state.ai_api_key:
        st.error(t["api_key_missing"])
    else:
        # 消耗授权次数（管理员除外）
        if not consume_usage(st.session_state.current_report_key):
            st.error(t["trial_ended"])
        else:
            # 准备提示词参数
            analyst_info = f"{analyst_name} ({analyst_title})" if analyst_title else analyst_name
            current_date = datetime.now().strftime("%Y-%m-%d")
            
            prompt = t["report_prompt"].format(
                product_name=product_name,
                product_description=product_desc,
                target_markets=", ".join(target_markets) if target_markets else "基于行业分析",
                target_users=target_users if target_users else "基于行业分析",
                CURRENT_DATE=current_date,
                ANALYST_INFO=analyst_info if analyst_info else "匿名分析师",
                channel_status=channel_status,
                channel_detail=channel_detail if channel_detail else "基于行业分析",
                brand_status=brand_status
            )
            
            # 生成报告
            with st.spinner(t["generating"]):
                try:
                    openai.api_key = st.session_state.ai_api_key
                    openai.base_url = st.session_state.ai_base_url
                    
                    response = openai.chat.completions.create(
                        model=st.session_state.ai_model_name,
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.3,
                        max_tokens=8000
                    )
                    
                    report_content = response.choices[0].message.content
                    
                    # 保存到session state
                    if lang == "zh":
                        st.session_state.report_content_zh = report_content
                    else:
                        st.session_state.report_content_en = report_content
                    
                    # 显示报告
                    st.subheader(t["report_title"])
                    st.markdown(report_content)
                    
                    # 下载功能（仅授权用户/管理员）
                    if is_premium_user(st.session_state.current_report_key):
                        st.subheader(t["download_section"])
                        # 生成Word文档
                        doc = Document()
                        markdown_to_docx(report_content, doc, lang)
                        
                        # 保存到BytesIO
                        buffer = BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        
                        # 下载按钮
                        st.download_button(
                            label=t["download_btn"],
                            data=buffer,
                            file_name=f"{product_name}_可行性分析报告_{current_date}.docx" if lang == "zh" else f"{product_name}_Feasibility_Report_{current_date}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    else:
                        st.info(t["no_license"])
                        
                except Exception as e:
                    st.error(f"{t['error_prefix']} {str(e)}")

# ================== 水印和安全CSS ==================
add_security_css(disable=is_premium_user(st.session_state.current_report_key))
add_dynamic_watermark(lang, hide=is_premium_user(st.session_state.current_report_key))

# ================== 页脚 ==================
st.markdown(f"""
<div style="text-align: center; margin-top: 50px; color: #666; font-size: 12px;">
    {t["footer"]}
</div>
""", unsafe_allow_html=True)
